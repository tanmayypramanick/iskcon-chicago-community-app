"""Build docs/seva-flows.docx — the seva reference, as an editable Word file.

Keeps docs/seva-flows.html and docs/seva-flows.docx saying the same thing.
Edit this, then run it from the repository root:

    python3 docs/build-seva-flows-docx.py

Needs python-docx.
"""

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

# The app's own palette, so the document matches the product it describes.
INK = RGBColor(0x3A, 0x34, 0x2B)
MUTED = RGBColor(0x76, 0x6D, 0x61)
POSTED = RGBColor(0x2B, 0x3A, 0x67)
REGISTRATION = RGBColor(0x1F, 0x6F, 0x67)
WEEKLY = RGBColor(0xB0, 0x70, 0x0D)
OPEN = RGBColor(0xC1, 0x44, 0x0E)

HEAD_FONT = "Georgia"
BODY_FONT = "Calibri"

doc = Document()

# Page setup and the default body style.
for section in doc.sections:
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)

normal = doc.styles["Normal"]
normal.font.name = BODY_FONT
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.15


def shade(cell, hex_colour):
    """Word has no cell-background property in python-docx; this is the XML for it."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_colour)
    tc_pr.append(shd)


def para(text="", size=10.5, colour=INK, bold=False, italic=False,
         space_before=0, space_after=8, font=BODY_FONT):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = colour
    run.bold = bold
    run.italic = italic
    return p


def heading(text, size=18, colour=INK, space_before=20, space_after=4):
    return para(text, size=size, colour=colour, bold=False,
                space_before=space_before, space_after=space_after,
                font=HEAD_FONT)


def eyebrow(text, colour=REGISTRATION):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.font.name = BODY_FONT
    run.font.size = Pt(8)
    run.font.color.rgb = colour
    run.bold = True
    return p


def table(headers, rows, widths=None, accent_col=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, text in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(text.upper())
        run.font.name = BODY_FONT
        run.font.size = Pt(8)
        run.bold = True
        run.font.color.rgb = MUTED
        shade(cell, "F5EFE3")

    for row in rows:
        cells = t.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(text)
            run.font.name = BODY_FONT
            run.font.size = Pt(9.5)
            # The third column of the outcome tables carries the verdict, and
            # its colour is information: green earns, red is refused.
            if accent_col is not None and i == accent_col:
                run.bold = True
                low = text.lower()
                if low.startswith("refused") or "refused" in low:
                    run.font.color.rgb = OPEN
                elif low.startswith("nothing") or low == "—":
                    run.font.color.rgb = MUTED
                else:
                    run.font.color.rgb = REGISTRATION
            else:
                run.font.color.rgb = INK

    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return t


def bullets(items, colour=REGISTRATION):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        run.font.name = BODY_FONT
        run.font.size = Pt(10)
        run.font.color.rgb = INK


# ---------------------------------------------------------------- title ----

eyebrow("ISKCON Chicago")
heading("How seva is settled", size=26, space_before=0, space_after=6)
para(
    "There are three kinds of seva and they are settled by different people in "
    "different ways. Nothing here is aspirational — every rule below is enforced "
    "by the database and proven by a test that runs whenever the schema is applied.",
    size=11, colour=MUTED, space_after=14,
)

# ---------------------------------------------------------------- flows ----

heading("The three flows")
para(
    "The one distinction everything else follows from: a posted seva is servable, "
    "a self-added seva is verifiable, and a weekly rota is neither — it runs by itself.",
    colour=MUTED, space_after=10,
)

flows = [
    (
        "Posted seva", "SERVABLE", POSTED,
        [
            ("Who posts it", "President, Tech Admin, Community Head or Volunteer."),
            ("Who serves it", "Any devotee joins an open seva; named devotees accept an invite-only one."),
            ("Who settles it", "Whoever posted it, or a Tech Admin or the President — they mark each place served, absent or excused, then complete the seva."),
            ("If a devotee can’t make it", "They step back before it starts. An open place simply reopens; an invited one raises a coverage request for whoever posted it."),
        ],
        "Not verifiable. The person who would verify it is the person who posted it, "
        "so marking somebody served is the verification.",
    ),
    (
        "Self-added seva", "VERIFIABLE", REGISTRATION,
        [
            ("Who adds it", "The devotee — either planned ahead, or logged after they have done it."),
            ("Who confirms it", "The Community Head, Tech Admin or President the devotee named. They cannot name themselves."),
            ("When", "Only after the seva has ended. Nothing is verified before it happens. Declining is allowed at any time."),
        ],
        "Not servable. There is no “mark served” step — not for the devotee, not for a "
        "Community Head, not for the President. Verifying it is what says it happened.",
    ),
    (
        "Weekly seva", "AUTOMATIC", WEEKLY,
        [
            ("Who shapes it", "Community Head, Tech Admin or President — the days, the places, and who stands on the rota."),
            ("How it settles", "It completes on its day and counts on completion alone. Nobody presses anything."),
            ("When somebody can’t go", "They ask for coverage. That names a substitute, opens the day, or asks another devotee."),
            ("Afterwards", "The devotee is asked whether they served it or missed it. Saying “missed” gives the credit back; ignoring it changes nothing. A cross puts the question away, which also changes nothing."),
        ],
        "No served or absent. Marking a rota settled nothing — the day still had nobody "
        "on it. Coverage is the mechanism that actually fills it.",
    ),
]

for name, verdict, colour, fields, note in flows:
    heading(name, size=14, colour=colour, space_before=14, space_after=2)
    para(verdict, size=8, colour=colour, bold=True, space_after=6)
    for label, value in fields:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        lab = p.add_run(label + "  ")
        lab.font.name = BODY_FONT
        lab.font.size = Pt(8.5)
        lab.bold = True
        lab.font.color.rgb = MUTED
        val = p.add_run(value)
        val.font.name = BODY_FONT
        val.font.size = Pt(10)
        val.font.color.rgb = INK
    para(note, size=9.5, colour=INK, italic=True, space_after=4)

# ------------------------------------------------------------ who may do ----

doc.add_page_break()
heading("Who may do what", space_before=0)
para(
    "Every one of these is checked on the server, so the answer is the same "
    "whatever the screen offers.",
    colour=MUTED, space_after=10,
)
table(
    ["Action", "Posted", "Self-added", "Weekly"],
    [
        ["Mark served / absent", "Poster, Tech, President", "Refused — nobody", "Refused — nobody"],
        ["Mark your own attendance", "Refused", "Refused", "Refused"],
        ["Verify it", "Nothing to verify", "The named member, Tech, President", "Nothing to verify"],
        ["Verify your own", "—", "Refused, even for the President", "—"],
        ["Complete it", "Poster, Tech, President", "Settled by verifying", "Automatic, on the day"],
        ["Step back before it starts", "The devotee holding the place", "Remove it instead — it is their own record", "Ask for coverage instead"],
        ["Say whether you served it", "The coordinator records it", "The named member verifies it", "The devotee, after it ends"],
        ["Arrange coverage", "—", "—", "The devotee asks; poster, Tech or President arranges"],
    ],
    widths=[1.5, 1.6, 1.9, 1.9],
)

# ------------------------------------------------------------ outcomes ------

heading("Every way a seva ends")
para(
    "A one-off act earns points when it is completed and verified and served — all "
    "three. A weekly act earns on completion alone. Being marked absent, excused, "
    "withdrawn or a no-show earns nothing, in either kind.",
    colour=MUTED, space_after=10,
)

para("Posted seva", size=10, bold=True, space_after=4)
table(
    ["What happened", "Result", "Earns"],
    [
        ["Devotee joined, coordinator marked them served, seva completed", "Counted", "Full hours"],
        ["Marked absent, or excused", "Not served", "Nothing"],
        ["Nobody answered for them", "Awaiting verification", "Nothing"],
        ["Some served, one marked absent", "Seva still completes", "Only those who served"],
        ["Everybody on it marked absent", "Seva becomes cancelled — it did not happen", "Nothing"],
        ["Marked served, then corrected to absent", "The correction is the last word", "Nothing, and points already given are taken back"],
        ["Arrived by QR scan, then marked served", "The scan is kept — a stronger record is never overwritten", "Full hours"],
        ["Devotee steps back from an OPEN seva before it starts", "The place reopens at once for anyone; whoever posted it is told", "Nothing for them"],
        ["Devotee steps back from an INVITED seva before it starts", "A coverage request opens for whoever posted it, who can open the day or ask somebody else", "Nothing for them"],
        ["Devotee tries to step back once it has started", "Refused — by then it is the coordinator’s to record", "Whatever the coordinator records"],
        ["Coordinator served a seva they posted themselves", "Refused — somebody else has to say so", "Nothing until another member records it"],
    ],
    widths=[2.6, 2.3, 2.0], accent_col=2,
)

para("Self-added seva", size=10, bold=True, space_after=4)
table(
    ["What happened", "Result", "Earns"],
    [
        ["Logged after the seva, the named member verified it", "Counted — verifying records that it was served", "Full hours"],
        ["Added, nobody has answered yet", "Exists only as a request — no seva record at all", "Nothing"],
        ["The member declined it", "Nothing is created, and nothing is left behind", "Nothing"],
        ["Planned ahead, member tries to verify before it happens", "Refused until the seva has ended", "Nothing yet"],
        ["Devotee tries to verify their own", "Refused", "Nothing"],
        ["Devotee tries to mark themselves served", "Refused — this flow has no attendance step", "Nothing"],
        ["Two registrations that overlap in time", "Both accepted; the clash is a warning, never a bar", "Both, if both are verified"],
    ],
    widths=[2.6, 2.3, 2.0], accent_col=2,
)

doc.add_page_break()
para("Weekly seva and coverage", size=10, bold=True, space_after=4)
table(
    ["What happened", "Result", "Earns"],
    [
        ["The day came and the occurrence completed", "Counted, with nobody pressing anything", "Full hours"],
        ["Anybody tries to mark served or absent", "Refused — ask for coverage instead", "Still counts on completion"],
        ["Devotee is asked afterwards and says “I served it”", "Recorded served — the credit the rota would have given anyway", "Full hours"],
        ["Devotee is asked afterwards and says “I missed it”", "Recorded absent, and whoever set the rota up is told the day went uncovered", "Nothing"],
        ["Devotee ignores the question", "Nothing changes — the prompt never costs anybody their hours", "Full hours"],
        ["Devotee puts the question away with the cross", "It stops being asked about that day. No answer is recorded on their behalf, and nobody is told", "Full hours"],
        ["Devotee says they cannot make one day", "A coverage request opens for that date", "Nothing for them that day"],
        ["Cannot make a run of days", "A date-range release, limited to the days they hold", "Nothing for them on those days"],
        ["Coordinator asks a specific devotee to cover", "They accept or decline; a decline reopens the day", "The substitute, for the day they take"],
        ["Coordinator opens the day to everyone", "The first devotee to take it fills it", "Whoever takes it"],
        ["The substitute suggests a different time", "The poster, Tech or President answers the counter-offer", "On the agreed time"],
        ["Nobody covers it", "The day stays open and visible in the coverage inbox", "Nothing"],
        ["A swap is accepted", "The standing rota is unchanged; the seva returns after the swap ends", "The substitute, only for the days covered"],
    ],
    widths=[2.6, 2.3, 2.0], accent_col=2,
)

# -------------------------------------------------------------- timing -----

heading("When things may happen")
table(
    ["Step", "Earliest it is allowed"],
    [
        ["Record attendance", "Once the seva has started — people are marked in as they arrive"],
        ["Complete a seva", "Once it has ended"],
        ["Verify a self-added seva", "Once it has ended"],
        ["Count towards hours and points", "Once it has ended, or been marked completed — a seva that finished early is not withheld"],
        ["Step back from a posted seva", "Any time before it starts, and not after"],
        ["Answer “did you serve your weekly seva?”", "Once it has ended. Answering is optional and never costs hours"],
        ["Put that question away", "Any time it is being asked, and only by the devotee it is asked of"],
        ["Correct attendance", "Any time. It reopens and recounts the Seva Mala week even after it has been frozen"],
    ],
    widths=[2.0, 4.9],
)
para(
    "Every one of these is decided in Chicago time, so the answer is the same for a "
    "devotee reading the app in Mayapur.",
    colour=MUTED,
)

# -------------------------------------------------------------- status -----

heading("Where this stands")
para(
    "Everything in the first list is live on the temple’s database and covered by a "
    "test. Everything in the second is known and not yet done.",
    colour=MUTED, space_after=10,
)

para("Working, and proven", size=11, colour=REGISTRATION, bold=True, space_after=4)
bullets([
    "The three flows are separated by the database, not by the screen",
    "Nobody records their own attendance, at any level",
    "Verification waits until the seva has ended",
    "Weekly seva cannot be marked served or absent",
    "A posted seva earns when the coordinator says it was served",
    "A self-added seva earns when the named member verifies it",
    "Absence removes points even after the week has been frozen",
    "Nothing is credited before it has happened",
    "Clash warnings inform and never block a registration",
    "A seva nobody served reads as cancelled, never completed",
    "Partial absence still completes, crediting only those who served",
    "A devotee can step back from a posted seva before it starts, and an invited "
    "place raises coverage rather than going quietly unfilled",
    "Weekly devotees are asked afterwards whether they served, and silence still counts",
    "Weekly coverage is walked end to end: asked and declined, asked again and "
    "accepted, opened to everyone, released as a run of days, a weekday the rota "
    "does not run on refused, and the same day released twice opening only one request",
    "The substitute is credited for the day they cover, and the devotee who "
    "released it is not",
])

para("Still to do", size=11, colour=OPEN, bold=True, space_before=8, space_after=4)
bullets([
    "Two coverage paths are still unproven: a range plan being superseded when the "
    "days are released again, and a substitute’s counter-offer of a different time",
    "The exported hours report still derives some rows from the registration rather "
    "than the seva record",
    "Roughly 35 tap targets across the seva screens are under 44pt",
    "“This service could not be found” covers three different situations, including a "
    "load that simply failed",
    "The seva list blames the filters when the fetch failed",
    "Seva history older than 180 days is outside the window the app reads",
])

para(
    "How to read a claim here. Each rule is enforced in the database and exercised by "
    "supabase/verification/seva_flow_matrix.sql, which drives posted, self-added and "
    "weekly seva through the real functions and checks what a devotee actually earns. "
    "Three of its expectations were wrong when it was written, and the database was "
    "right each time.",
    size=9, colour=MUTED, space_before=14,
)

doc.save("docs/seva-flows.docx")
print("saved docs/seva-flows.docx")
